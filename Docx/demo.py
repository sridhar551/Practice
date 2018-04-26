import docx
#import readDocx

doc = docx.Document('/media/sree/E87E40977E406084/Sridhar/Resumes/Sanjay/1.docx')
print(len(doc.paragraphs))
print(doc)

#title = doc.paragraphs[0].text

#print(readDocx.getText('/media/sree/E87E40977E406084/Sridhar/Resumes/Sanjay/1.docx'))